# your_package/document_processor/docx_handler.py
import logging, os, tempfile, asyncio
from typing import Any, Dict, List, Set
from pathlib import Path
from docx import Document

from .utils import clean_text
from .ocr import convert_images_to_text_batch, convert_pdf_to_markdown_with_html_reference
from .config import is_image_text_enabled
from .html_reprocessor import clean_html_file
logger = logging.getLogger("document-processor")

try:
    from pdf2image import convert_from_path
    PDF2IMAGE_AVAILABLE = True
except Exception:
    PDF2IMAGE_AVAILABLE = False

try:
    from docx2pdf import convert as docx_to_pdf_convert
    DOCX2PDF_AVAILABLE = True
except Exception:
    DOCX2PDF_AVAILABLE = False

def _has_page_break_element(element) -> bool:
    try:
        nsmap = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        if element.findall('.//w:br[@w:type="page"]', nsmap): return True
        if element.findall('.//w:lastRenderedPageBreak', nsmap): return True
        return False
    except Exception:
        return False

def _paragraph_has_page_break(paragraph) -> bool:
    try:
        for run in paragraph.runs:
            if run.element.findall('.//w:br[@w:type="page"]',
                                   {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
                return True
        if hasattr(paragraph, '_element'):
            return _has_page_break_element(paragraph._element)
        return False
    except Exception:
        return False

def _extract_paragraph_text(para_element, doc) -> str:
    try:
        text = ""
        nsmap = doc.element.nsmap if hasattr(doc.element, 'nsmap') else {}
        for run in para_element.findall('.//w:r', nsmap):
            for t in run.findall('.//w:t', nsmap):
                if t.text: text += t.text
            for drawing in run.findall('.//w:drawing', nsmap):
                text += " [이미지] "
                try:
                    for desc in drawing.findall('.//wp:docPr', nsmap):
                        if desc.get('descr'): text += f"[설명: {desc.get('descr')}] "
                        elif desc.get('name'): text += f"[이름: {desc.get('name')}] "
                except:
                    pass
        return text
    except Exception:
        try:
            ts = para_element.findall('.//w:t', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
            return ''.join([e.text or '' for e in ts])
        except:
            return para_element.text or ""

def _extract_table_text_xml(table_element) -> str:
    try:
        ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        out = ""
        for row in table_element.findall('.//w:tr', ns):
            row_text = []
            for cell in row.findall('.//w:tc', ns):
                cell_text = ""
                for t in cell.findall('.//w:t', ns):
                    if t.text: cell_text += t.text
                row_text.append(cell_text.strip())
            if any(c.strip() for c in row_text):
                out += " | ".join(row_text) + "\n"
        return out
    except Exception:
        return ""

def _extract_simple_table_text(table) -> str:
    try:
        out = ""
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                row_text.append(cell.text.strip())
            if any(c.strip() for c in row_text):
                out += " | ".join(row_text) + "\n"
        return out
    except Exception:
        return ""

def _is_similar_table_text(t1: str, t2: str, threshold: float = 0.8) -> bool:
    import re
    if not t1 or not t2: return False
    a = re.sub(r'\s+', ' ', t1.strip()); b = re.sub(r'\s+', ' ', t2.strip())
    if a == b: return True
    l1, l2 = len(a), len(b)
    if min(l1, l2)/max(l1, l2) < 0.5: return False
    s = a if l1 < l2 else b
    L = b if l1 < l2 else a
    ratio = len(set(s.split()) & set(L.split())) / len(set(s.split()))
    return ratio >= threshold

async def convert_docx_to_pdf_libreoffice(file_path: str) -> str:
    """LibreOffice로 DOCX를 PDF로 변환"""
    try:
        import subprocess
        with tempfile.TemporaryDirectory() as temp_dir:
            cmd = [
                'libreoffice', '--headless', '--convert-to', 'pdf', 
                '--outdir', temp_dir, file_path
            ]
            result = subprocess.run(cmd, check=True, capture_output=True, text=True)
            
            docx_name = Path(file_path).stem
            pdf_path = os.path.join(temp_dir, f"{docx_name}.pdf")
            
            if not os.path.exists(pdf_path):
                raise FileNotFoundError(f"PDF conversion failed: {pdf_path}")
            
            # 임시 파일로 복사
            with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as temp_pdf:
                with open(pdf_path, 'rb') as f:
                    temp_pdf.write(f.read())
                logger.info(f"DOCX → PDF 변환 완료: {temp_pdf.name}")
                return temp_pdf.name
                
    except subprocess.CalledProcessError as e:
        logger.error(f"LibreOffice PDF conversion failed: {e.stderr}")
        raise
    except Exception as e:
        logger.error(f"PDF conversion error: {e}")
        raise

async def convert_docx_to_images(file_path: str) -> List[str]:
    temp_files: List[str] = []
    try:
        if DOCX2PDF_AVAILABLE and PDF2IMAGE_AVAILABLE:
            with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as tp:
                docx_to_pdf_convert(file_path, tp.name)
                images = convert_from_path(tp.name, dpi=300)
                for im in images:
                    with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as ti:
                        im.save(ti.name, 'PNG')
                        temp_files.append(ti.name)
                os.unlink(tp.name)
            return temp_files
        # LibreOffice CLI fallback
        elif PDF2IMAGE_AVAILABLE:
            import subprocess
            with tempfile.TemporaryDirectory() as td:
                try:
                    subprocess.run(['libreoffice','--headless','--convert-to','pdf','--outdir', td, file_path],
                                   check=True, capture_output=True)
                    pdf_path = os.path.join(td, Path(file_path).stem + '.pdf')
                    if os.path.exists(pdf_path):
                        images = convert_from_path(pdf_path, dpi=300)
                        for im in images:
                            with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as ti:
                                im.save(ti.name, 'PNG')
                                temp_files.append(ti.name)
                        return temp_files
                except Exception:
                    pass
        return []
    except Exception:
        for p in temp_files:
            try: os.unlink(p)
            except: pass
        return []

async def convert_docx_to_html_text(file_path: str) -> str:
    """DOCX를 HTML로 변환 후 정리된 HTML 반환"""
    try:
        import subprocess
        
        with tempfile.TemporaryDirectory() as temp_dir:
            cmd = [
                'libreoffice', '--headless', '--convert-to', 'html',
                '--outdir', temp_dir, file_path
            ]
            subprocess.run(cmd, check=True, capture_output=True)
            
            docx_name = Path(file_path).stem
            html_path = os.path.join(temp_dir, f"{docx_name}.html")
            
            if not os.path.exists(html_path):
                raise FileNotFoundError("HTML conversion failed")
            
            # HTML 파일 읽기
            with open(html_path, 'r', encoding='utf-8') as f:
                html_content = f.read()
            
            # HTML 정리해서 반환
            cleaned_html = clean_html_file(html_content)
            return cleaned_html
            
    except Exception as e:
        logger.error(f"HTML conversion failed: {e}")
        raise

# 1. 기본 텍스트 추출 (python-docx 직접 사용)
async def extract_text_from_docx_fallback(file_path: str) -> str:
    """python-docx로 직접 텍스트 추출"""
    doc = Document(file_path)
    text = ""
    processed: Set[str] = set()
    current_page = 1
    try:
        for element in doc.element.body:
            if element.tag.endswith('p'):
                if _has_page_break_element(element):
                    current_page += 1
                    text += f"\n<페이지 번호> {current_page} </페이지 번호>\n"
                para_text = _extract_paragraph_text(element, doc)
                if para_text.strip():
                    text += para_text + "\n"
            elif element.tag.endswith('tbl'):
                t = _extract_table_text_xml(element)
                if t.strip():
                    text += "\n=== 표 ===\n" + t + "\n=== 표 끝 ===\n\n"
                    processed.add(t)
    except Exception:
        text = ""
        current_page = 1
        for p in doc.paragraphs:
            if _paragraph_has_page_break(p):
                current_page += 1
                text += f"\n<페이지 번호> {current_page} </페이지 번호>n"
            if p.text.strip():
                text += p.text + "\n"

    for i, table in enumerate(doc.tables):
        t = _extract_simple_table_text(table)
        if t.strip() and not any(_is_similar_table_text(t, x) for x in processed):
            text += f"\n=== 표 {i+1} ===\n{t}\n=== 표 끝 ===\n\n"
            processed.add(t)

    if current_page > 1 and not text.startswith("<페이지 번호>  1 </페이지 번호>"):
        text = f"<페이지 번호>  1 </페이지 번호>n{text}"
    return clean_text(text)

# 2. HTML 변환 방식
async def extract_text_from_docx_fallback_html(file_path: str) -> str:
    """DOCX를 HTML로 변환 후 clean_html_file로 가공"""
    try:
        html_content = await convert_docx_to_html_text(file_path)
        logger.info("DOCX → HTML 변환 및 정리 완료")
        return html_content
    except Exception as e:
        logger.warning(f"HTML 변환 실패, 기존 fallback 사용: {e}")
        return await extract_text_from_docx_fallback(file_path)

# 3. 이미지 OCR 방식
async def extract_text_from_docx_via_ocr(file_path: str, current_config: Dict[str, Any]) -> str:
    """DOCX → 이미지 → OCR 처리"""
    if not is_image_text_enabled(current_config, True):
        logger.warning("DOCX: OCR requested but not enabled, falling back to text extraction")
        return await extract_text_from_docx_fallback(file_path)
    
    images = await convert_docx_to_images(file_path)
    if not images:
        logger.warning("DOCX: Image conversion failed, falling back to text extraction")
        return await extract_text_from_docx_fallback(file_path)
    
    try:
        batch_size = current_config.get('batch_size', 1)
        page_texts = await convert_images_to_text_batch(images, current_config, batch_size)
        all_text = ""
        for i, t in enumerate(page_texts):
            if not str(t).startswith("[이미지 파일:"):
                all_text += f"\n<페이지 번호> {i+1} (OCR) </페이지 번호>n{t}\n"
        
        if all_text.strip():
            logger.info(f"DOCX: OCR processing completed for {len(page_texts)} pages")
            return clean_text(all_text)
        else:
            logger.warning("DOCX: OCR failed, falling back to text extraction")
            return await extract_text_from_docx_fallback(file_path)
    finally:
        for p in images:
            try: os.unlink(p)
            except: pass

# 4. HTML+PDF OCR 복합 방식
async def extract_text_from_docx_via_html_pdf_ocr(file_path: str, current_config: Dict[str, Any]) -> str:
    """DOCX를 HTML(텍스트) + PDF(이미지)로 변환 후 OCR로 마크다운 생성"""
    if not is_image_text_enabled(current_config, True):
        logger.warning("DOCX: HTML+PDF OCR requested but not enabled, falling back to HTML extraction")
        return await extract_text_from_docx_fallback_html(file_path)
    
    pdf_path = None
    try:
        # 1. DOCX → HTML 텍스트 추출 (reference용)
        html_reference_text = await convert_docx_to_html_text(file_path)
        logger.info("DOCX → HTML 텍스트 추출 완료")
        
        # 2. DOCX → PDF 변환
        pdf_path = await convert_docx_to_pdf_libreoffice(file_path)
        logger.info("DOCX → PDF 변환 완료")
        
        # 3. PDF → 마크다운 (HTML 텍스트를 reference로 사용)
        markdown_result = await convert_pdf_to_markdown_with_html_reference(
            pdf_path, html_reference_text, current_config
        )
        
        if markdown_result and not markdown_result.startswith("["):
            logger.info("DOCX: HTML+PDF OCR processing completed")
            return clean_text(markdown_result)
        else:
            logger.warning("DOCX: HTML+PDF OCR failed, falling back to HTML extraction")
            return await extract_text_from_docx_fallback_html(file_path)
            
    except Exception as e:
        logger.error(f"DOCX: HTML+PDF OCR processing failed: {e}, falling back to HTML extraction")
        return await extract_text_from_docx_fallback_html(file_path)
    finally:
        if pdf_path and os.path.exists(pdf_path):
            try:
                os.unlink(pdf_path)
            except:
                pass

async def _extract_docx_default(file_path: str, current_config: Dict[str, Any]) -> str:
    """기존 DOCX 처리 로직 (자동 선택)"""
    provider = current_config.get('provider', 'no_model')
    
    if provider == 'no_model':
        return await extract_text_from_docx_fallback_html(file_path)  # HTML 방식 사용
    
    # 1순위: HTML+PDF OCR 방식
    try:
        return await extract_text_from_docx_fallback_html(file_path)  # HTML 방식 사용
    except Exception as e:
        logger.warning(f"HTML+PDF OCR failed, falling back to image OCR: {e}")
        # 2순위로 폴백
        return await extract_text_from_docx_via_ocr(file_path, current_config)

async def extract_text_from_docx(file_path: str, current_config: Dict[str, Any], process_type: str = "default") -> str:
    """DOCX 텍스트 추출 메인 함수"""
    provider = current_config.get('provider', 'no_model')
    logger.info(f"Real-time DOCX processing with provider: {provider}, process_type: {process_type}")
    
    if process_type == "text":
        # 1. 기본 텍스트 추출 (python-docx 직접 사용)
        logger.info("DOCX basic text extraction processing requested")
        return await extract_text_from_docx_fallback(file_path)
    
    elif process_type == "html":
        # 2. HTML 변환 방식
        logger.info("DOCX HTML processing requested")
        return await extract_text_from_docx_fallback_html(file_path)
    
    elif process_type == "ocr":
        # 3. 이미지 OCR 방식
        logger.info("DOCX OCR processing requested")
        return await extract_text_from_docx_via_ocr(file_path, current_config)
    
    elif process_type == "html_pdf_ocr":
        # 4. HTML+PDF OCR 복합 방식
        logger.info("DOCX HTML+PDF OCR processing requested")
        return await extract_text_from_docx_via_html_pdf_ocr(file_path, current_config)
    
    else:  # process_type == "default"
        # 기존 자동 선택 로직 유지
        return await _extract_docx_default(file_path, current_config)